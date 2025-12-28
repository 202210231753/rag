import os
import pandas as pd
from docx import Document
from app.rag.parsers.office_parser import WordParser, ExcelParser

def create_dummy_files():
    """创建临时的测试文件"""
    print("📄 Creating dummy Word file...")
    doc = Document()
    doc.add_heading('Test Document Title', 0)
    doc.add_paragraph('This is a normal paragraph.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content of section 1.')
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'City'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Alice'
    row_cells[1].text = '30'
    row_cells[2].text = 'New York'
    
    doc.save('test_doc.docx')
    
    print("📊 Creating dummy Excel file...")
    df = pd.DataFrame({
        'Product': ['Apple', 'Banana', 'Orange'],
        'Price': [1.2, 0.5, 0.8],
        'Stock': [100, 200, 150]
    })
    df.to_excel('test_sheet.xlsx', index=False)

def test_parsers():
    """测试解析器"""
    try:
        # 1. Test Word Parser
        print("\n🧪 Testing WordParser...")
        word_parser = WordParser()
        result = word_parser.parse('test_doc.docx')
        print("--- Word Parse Result ---")
        print(result['content'][:500]) # Print first 500 chars
        
        if "Test Document Title" in result['content'] and "| Alice |" in result['content']:
            print("✅ WordParser Test Passed!")
        else:
            print("❌ WordParser Test Failed!")

        # 2. Test Excel Parser
        print("\n🧪 Testing ExcelParser...")
        excel_parser = ExcelParser()
        result = excel_parser.parse('test_sheet.xlsx')
        print("--- Excel Parse Result ---")
        print(result['content'][:500])
        
        if "Apple" in result['content'] and "| 1.2 |" in result['content']:
            print("✅ ExcelParser Test Passed!")
        else:
            print("❌ ExcelParser Test Failed!")
            
    except Exception as e:
        print(f"❌ Error during testing: {e}")
    finally:
        # Cleanup
        if os.path.exists('test_doc.docx'):
            os.remove('test_doc.docx')
        if os.path.exists('test_sheet.xlsx'):
            os.remove('test_sheet.xlsx')

if __name__ == "__main__":
    create_dummy_files()
    test_parsers()
