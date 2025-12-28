import pandas as pd
from docx import Document as DocxDocument
from app.rag.parsers.base import BaseParser
from typing import Dict, Any
import os

class WordParser(BaseParser):
    """Word (.docx) 解析器"""
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        doc = DocxDocument(file_path)
        full_text = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                # 简单的 Markdown 转换：根据样式判断标题
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        prefix = '#' * level_num
                        full_text.append(f"{prefix} {para.text}")
                    except:
                        full_text.append(f"**{para.text}**")
                else:
                    full_text.append(para.text)
        
        # 提取表格 (简单转换为 Markdown 表格)
        for table in doc.tables:
            # 将表格转为 list of lists
            data = []
            keys = None
            for i, row in enumerate(table.rows):
                text = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    keys = text
                else:
                    data.append(text)
            
            if keys:
                try:
                    df = pd.DataFrame(data, columns=keys)
                    markdown_table = df.to_markdown(index=False)
                    full_text.append("\n" + markdown_table + "\n")
                except Exception as e:
                    # 降级处理：如果 pandas 转换失败，直接拼文本
                    print(f"Table parsing failed: {e}")
                    pass

        return {
            "content": "\n\n".join(full_text),
            "images": [], # python-docx 提取图片比较麻烦，暂略
            "meta": {"filename": os.path.basename(file_path)}
        }

class ExcelParser(BaseParser):
    """Excel (.xlsx, .csv) 解析器"""
    
    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        content_parts = []
        
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
                content_parts.append(f"### Sheet: CSV Data")
                content_parts.append(df.to_markdown(index=False))
            else:
                # 读取所有 Sheet
                xls = pd.ExcelFile(file_path)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    content_parts.append(f"### Sheet: {sheet_name}")
                    # 转换为 Markdown 表格
                    content_parts.append(df.to_markdown(index=False))
                    content_parts.append("\n")
                    
        except Exception as e:
            raise ValueError(f"Failed to parse Excel file: {e}")

        return {
            "content": "\n\n".join(content_parts),
            "images": [],
            "meta": {"filename": os.path.basename(file_path)}
        }
